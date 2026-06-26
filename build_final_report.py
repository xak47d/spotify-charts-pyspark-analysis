from __future__ import annotations

import json
import shutil
import tempfile
from zipfile import ZIP_DEFLATED, ZipFile
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from lxml import etree


ROOT = Path(__file__).resolve().parent
REPORT_DIR = ROOT / "report"
ASSETS = REPORT_DIR / "assets"
OUTPUT = REPORT_DIR / "Proyecto_Final_Equipo51.docx"
METRICS_PATH = REPORT_DIR / "final_metrics.json"
TEMPLATE = ROOT.parent.parent / "Plantilla-Tareas-TEC.docx"

BLACK = "000000"
BLUE = BLACK
DARK_BLUE = BLACK
INK = BLACK
MUTED = BLACK
LIGHT = "F4F6F9"
WHITE = "FFFFFF"

TEAM = [
    "Fernando Arango Gaviria — A01797660",
    "Jose Luis Armenta Mandujano — A01796933",
    "Demenard Gardy Armand — A01797139",
    "Ricardo Ismael Vega Aguilar — A01796617",
]

NOTEBOOK_LINK = "[ENLACE AL NOTEBOOK POR REEMPLAZAR DESPUÉS DE LA CARGA]"
VIDEO_LINK = "[ENLACE AL VIDEO POR REEMPLAZAR DESPUÉS DE LA CARGA]"
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
WP_NS = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
NSMAP = {"w": W_NS, "wp": WP_NS}


def set_font(run, size=12, bold=None, italic=None, color=INK, name="Arial"):
    # User requirement: the entire Word document must use Arial 12.
    size = 12
    name = "Arial"
    run.font.name = name
    run._element.get_or_add_rPr()
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def paragraph(doc, text="", *, bold=False, italic=False, size=12, color=INK,
              align=WD_ALIGN_PARAGRAPH.JUSTIFY, before=0, after=6, line=1.15):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    set_font(p.add_run(text), size=size, bold=bold, italic=italic, color=color)
    return p


def heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    set_font(p.add_run(text), size=12, bold=True,
             color=BLUE if level < 3 else DARK_BLUE)
    return p


def page_break(doc):
    doc.add_page_break()


def field(paragraph, instruction):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, end])
    set_font(run, size=9, color=MUTED)


def configure(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(12)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for level, size, before, after, color in [
        (1, 12, 14, 7, BLUE),
        (2, 12, 9, 4, BLUE),
        (3, 12, 6, 3, DARK_BLUE),
    ]:
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    # User requirement: no visible header or footer in the final report.
    for container in (section.header, section.footer):
        for p in container.paragraphs:
            p.clear()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)


def set_paragraph_text(p, text="", *, bold=False, italic=False, color=INK,
                       align=WD_ALIGN_PARAGRAPH.CENTER, before=None, after=None, line=1.15):
    p.clear()
    p.alignment = align
    if before is not None:
        p.paragraph_format.space_before = Pt(before)
    if after is not None:
        p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    run = p.add_run(text)
    set_font(run, bold=bold, italic=italic, color=color)
    return p


def setup_template_cover(doc):
    """Reuse the TEC assignment template cover and replace its editable text."""
    if len(doc.paragraphs) < 19:
        raise ValueError("La plantilla TEC no tiene la estructura esperada.")

    for shape in doc.inline_shapes:
        doc_pr = shape._inline.docPr
        doc_pr.set("descr", "Logotipo del Tecnológico de Monterrey")
        doc_pr.set("title", "Logotipo Tecnológico de Monterrey")

    for idx in [5, 6, 10, 11, 12, 15, 16, 17]:
        doc.paragraphs[idx].clear()
        doc.paragraphs[idx].paragraph_format.space_before = Pt(0)
        doc.paragraphs[idx].paragraph_format.space_after = Pt(0)
        doc.paragraphs[idx].paragraph_format.line_spacing = 1.0

    set_paragraph_text(
        doc.paragraphs[3],
        "Escuela de Ingeniería y Ciencias",
        bold=True,
        color=INK,
        after=7,
    )
    set_paragraph_text(
        doc.paragraphs[4],
        "Maestría en Inteligencia Artificial Aplicada",
        bold=True,
        color=INK,
        after=28,
    )
    set_paragraph_text(doc.paragraphs[7], "PRESENTA:", bold=True, after=12)

    authors = doc.paragraphs[8]
    authors.clear()
    authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
    authors.paragraph_format.line_spacing = 1.15
    authors.paragraph_format.space_after = Pt(8)
    for idx, member in enumerate(TEAM):
        run = authors.add_run(member)
        set_font(run, bold=True, color=INK)
        if idx != len(TEAM) - 1:
            run.add_break()

    set_paragraph_text(doc.paragraphs[9], "Equipo 51", bold=True, after=28)
    set_paragraph_text(doc.paragraphs[13], "ACTIVIDAD:", bold=True, after=12)
    set_paragraph_text(
        doc.paragraphs[14],
        "Clasificación de la visibilidad competitiva de artistas en Spotify Charts mediante aprendizaje distribuido",
        bold=True,
        color=INK,
        after=8,
        line=1.1,
    )
    set_paragraph_text(doc.paragraphs[18], "Junio de 2026", bold=True, before=16)


def enforce_arial_12(path: Path):
    """Normalize inherited/template OOXML runs and styles to Arial 12 pt."""
    with tempfile.TemporaryDirectory() as tmpdir:
        tmp = Path(tmpdir) / "normalized.docx"
        with ZipFile(path, "r") as zin, ZipFile(tmp, "w", ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename.startswith("word/") and item.filename.endswith(".xml"):
                    try:
                        root = etree.fromstring(data)
                    except etree.XMLSyntaxError:
                        zout.writestr(item, data)
                        continue

                    changed = False
                    for rpr in root.xpath(".//w:rPr", namespaces=NSMAP):
                        rfonts = rpr.find("w:rFonts", namespaces=NSMAP)
                        if rfonts is None:
                            rfonts = etree.SubElement(rpr, f"{{{W_NS}}}rFonts")
                        for attr in [
                            "ascii", "hAnsi", "eastAsia", "cs",
                            "asciiTheme", "hAnsiTheme", "eastAsiaTheme", "cstheme",
                        ]:
                            q_attr = f"{{{W_NS}}}{attr}"
                            if attr.endswith("Theme") or attr == "cstheme":
                                if q_attr in rfonts.attrib:
                                    del rfonts.attrib[q_attr]
                            else:
                                rfonts.set(q_attr, "Arial")

                        for tag in ["sz", "szCs"]:
                            node = rpr.find(f"w:{tag}", namespaces=NSMAP)
                            if node is None:
                                node = etree.SubElement(rpr, f"{{{W_NS}}}{tag}")
                            node.set(f"{{{W_NS}}}val", "24")
                        color = rpr.find("w:color", namespaces=NSMAP)
                        if color is None:
                            color = etree.SubElement(rpr, f"{{{W_NS}}}color")
                        color.set(f"{{{W_NS}}}val", BLACK)
                        changed = True

                    if item.filename.endswith("styles.xml"):
                        for rpr_default in root.xpath(".//w:rPrDefault/w:rPr", namespaces=NSMAP):
                            rfonts = rpr_default.find("w:rFonts", namespaces=NSMAP)
                            if rfonts is None:
                                rfonts = etree.SubElement(rpr_default, f"{{{W_NS}}}rFonts")
                            for attr in [
                                "ascii", "hAnsi", "eastAsia", "cs",
                                "asciiTheme", "hAnsiTheme", "eastAsiaTheme", "cstheme",
                            ]:
                                q_attr = f"{{{W_NS}}}{attr}"
                                if attr.endswith("Theme") or attr == "cstheme":
                                    if q_attr in rfonts.attrib:
                                        del rfonts.attrib[q_attr]
                                else:
                                    rfonts.set(q_attr, "Arial")
                            for tag in ["sz", "szCs"]:
                                node = rpr_default.find(f"w:{tag}", namespaces=NSMAP)
                                if node is None:
                                    node = etree.SubElement(rpr_default, f"{{{W_NS}}}{tag}")
                                node.set(f"{{{W_NS}}}val", "24")
                            color = rpr_default.find("w:color", namespaces=NSMAP)
                            if color is None:
                                color = etree.SubElement(rpr_default, f"{{{W_NS}}}color")
                            color.set(f"{{{W_NS}}}val", BLACK)
                            changed = True

                    for doc_pr in root.xpath(".//wp:docPr", namespaces=NSMAP):
                        if not doc_pr.get("descr") and not doc_pr.get("title"):
                            doc_pr.set("descr", "Logotipo del Tecnológico de Monterrey")
                            doc_pr.set("title", "Logotipo Tecnológico de Monterrey")
                            changed = True

                    if changed:
                        data = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
                zout.writestr(item, data)
        shutil.move(str(tmp), path)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def margins(cell, top=80, bottom=80, start=120, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in [("top", top), ("bottom", bottom), ("start", start), ("end", end)]:
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    total_dxa = 9360
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total_dxa))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    widths_dxa = [int(w * 1440) for w in widths]
    for value in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(value))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            margins(cell)


def set_table_borders(table, color="D9E2EC", size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        node = borders.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def table(doc, headers, rows, widths, font_size=9.5):
    t = doc.add_table(rows=1, cols=len(headers))
    if "Table Grid" in [s.name for s in doc.styles]:
        t.style = "Table Grid"
    set_table_borders(t)
    tr_pr = t.rows[0]._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)
    for idx, label in enumerate(headers):
        cell = t.rows[0].cells[idx]
        cell.text = ""
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        shade(cell, LIGHT)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        set_font(p.add_run(str(label)), size=font_size, bold=True, color=DARK_BLUE)
    for row in rows:
        cells = t.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = ""
            cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cells[idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.1
            set_font(p.add_run(str(value)), size=font_size, color=INK)
    set_table_geometry(t, widths)
    paragraph(doc, "", after=3, line=1.0)
    return t


def figure(doc, filename, caption, width=5.7):
    path = ASSETS / filename
    if not path.exists():
        raise FileNotFoundError(path)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3)
    shape = p.add_run().add_picture(str(path), width=Inches(width))
    doc_pr = shape._inline.docPr
    doc_pr.set("descr", caption)
    doc_pr.set("title", caption.split(". ", 1)[0])
    cap = paragraph(
        doc,
        caption,
        italic=True,
        size=9,
        color=MUTED,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=7,
        line=1.0,
    )
    cap.paragraph_format.keep_with_next = False


def bullets(doc, items):
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.194)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.15
        set_font(p.add_run(f"• {item}"), size=12, color=INK)


def metric(metrics, name, kind="mean"):
    return metrics["random_forest"][f"metrics_{kind}"][name]


def build():
    if not METRICS_PATH.exists():
        raise FileNotFoundError("Ejecute ProyectoFinal_Equipo51.ipynb antes de construir el reporte.")

    metrics = json.loads(METRICS_PATH.read_text(encoding="utf-8"))
    strata = pd.read_csv(REPORT_DIR / "tabla_estratos.csv")
    folds = pd.read_csv(REPORT_DIR / "tabla_metricas_folds.csv")
    baseline = pd.read_csv(REPORT_DIR / "tabla_modelo_vs_linea_base.csv")
    per_class = pd.read_csv(REPORT_DIR / "tabla_metricas_por_clase.csv")
    tuning = pd.read_csv(REPORT_DIR / "tabla_hiperparametros.csv")
    importance = pd.read_csv(REPORT_DIR / "tabla_importancia_variables.csv")
    kmeans = pd.read_csv(REPORT_DIR / "tabla_kmeans.csv")
    centers = pd.read_csv(REPORT_DIR / "tabla_centroides_clusters.csv")

    if not TEMPLATE.exists():
        raise FileNotFoundError(f"No se encontró la plantilla TEC: {TEMPLATE}")

    doc = Document(TEMPLATE)
    configure(doc)
    setup_template_cover(doc)

    # 2 — resumen
    page_break(doc)
    heading(doc, "Resumen", 1)
    paragraph(
        doc,
        f"Este proyecto presenta un flujo integral de aprendizaje automático distribuido para clasificar observaciones diarias de artistas en Spotify Charts como Top o Estable. La población original contiene {metrics['dataset']['original_records']:,} registros con cobertura de {metrics['dataset']['date_min']} a {metrics['dataset']['date_max']}. Después de excluir Venezuela por una anomalía de carga y establecer una ventana con historial suficiente, la población evaluable quedó en {metrics['dataset']['filtered_records']:,} observaciones.",
    )
    paragraph(
        doc,
        f"La población se caracterizó por antigüedad en chart, alcance geográfico y posición competitiva. Con esas variables se construyeron ocho estratos y una muestra proporcional sin reemplazo de {metrics['sample']['actual']:,} registros. El modelo principal fue Random Forest, con ajuste de número de árboles y profundidad, seguido por validación cruzada estratificada de cinco pliegues. La configuración seleccionada utilizó {metrics['random_forest']['selected_numTrees']} árboles y profundidad máxima {metrics['random_forest']['selected_maxDepth']}.",
    )
    paragraph(
        doc,
        f"El desempeño promedio alcanzó AUC-ROC de {metric(metrics, 'auc_roc'):.4f}, F1 ponderado de {metric(metrics, 'f1_weighted'):.4f} y accuracy de {metric(metrics, 'accuracy'):.4f}. La desviación estándar de AUC-ROC fue {metric(metrics, 'auc_roc', 'std'):.4f}, lo que permite valorar la estabilidad entre particiones. KMeans se incorporó como análisis complementario y seleccionó k={metrics['kmeans']['best_k']} con Silhouette de {metrics['kmeans']['silhouette']:.4f}.",
    )
    paragraph(
        doc,
        f"La línea base que siempre predice la clase mayoritaria alcanzó accuracy de {metrics['random_forest']['baseline_mean']['accuracy']:.4f}, superior por {abs(metrics['random_forest']['model_minus_baseline']['accuracy']):.4f} a la del bosque. Sin embargo, esa línea base tiene recall cero para la clase Top. Random Forest recuperó Top con recall de {metrics['random_forest']['per_class']['Top']['recall']:.4f} y mejoró el F1 ponderado en {metrics['random_forest']['model_minus_baseline']['f1_weighted']:.4f}. Por ello, la aportación no consiste en superar la línea base en todas las métricas, sino en obtener separación útil y recuperar la clase minoritaria sin utilizar directamente el ranking como predictor.",
    )
    paragraph(doc, "Palabras clave: Big Data, PySpark, Spotify Charts, Random Forest, validación cruzada, muestreo estratificado.", italic=True, color=MUTED)

    # 3 — introducción
    page_break(doc)
    heading(doc, "1. Introducción", 1)
    paragraph(
        doc,
        "Los servicios de streaming transformaron la circulación de la música en un proceso observable a escala mundial. Los charts diarios sintetizan millones de interacciones y permiten analizar la persistencia de artistas, la amplitud territorial de su presencia y su nivel competitivo. Sin embargo, el volumen y la heterogeneidad de estas observaciones hacen insuficiente un flujo de análisis exclusivamente en memoria.",
    )
    paragraph(
        doc,
        "El problema de investigación consiste en determinar si las características históricas y geográficas de una observación permiten reconocer si un artista se encuentra en la franja Top del chart o en una posición Estable. La pregunta no pretende pronosticar el ranking futuro ni establecer causalidad; busca identificar patrones explicativos dentro de la población observada.",
    )
    heading(doc, "1.1 Objetivo general", 2)
    paragraph(
        doc,
        "Construir y evaluar un modelo de aprendizaje distribuido capaz de clasificar la visibilidad competitiva de artistas en Spotify Charts, preservando la estructura de una población de gran volumen y documentando todas las etapas del proceso.",
    )
    heading(doc, "1.2 Objetivos específicos", 2)
    bullets(doc, [
        "Auditar y delimitar una población confiable a partir del archivo charts_artists_daily.csv.",
        "Construir una muestra representativa mediante estratificación proporcional.",
        "Prevenir fuga de información y comparar configuraciones de Random Forest.",
        "Medir desempeño y variabilidad mediante cinco folds estratificados.",
        "Complementar la clasificación con perfiles no supervisados de KMeans.",
    ])

    # 4 — solución y datos
    page_break(doc)
    heading(doc, "2. Propuesta de solución", 1)
    paragraph(
        doc,
        "La solución se implementó en PySpark para distribuir lectura, agregaciones, muestreo y entrenamiento. Pandas y las bibliotecas de visualización se utilizaron únicamente después de reducir la información a tablas agregadas. El flujo se organizó en siete etapas: auditoría, caracterización, muestreo, preparación, ajuste, validación cruzada y análisis complementario.",
    )
    heading(doc, "2.1 Fuente y calidad de datos", 2)
    table(doc, ["Elemento", "Resultado"], [
        ("Archivo", "charts_artists_daily.csv"),
        ("Registros originales", f"{metrics['dataset']['original_records']:,}"),
        ("Cobertura", f"{metrics['dataset']['date_min']} a {metrics['dataset']['date_max']}"),
        ("Población evaluable", f"{metrics['dataset']['filtered_records']:,}"),
        ("Exclusión", "Venezuela (VE), por patrón de carga anómalo"),
        ("Inicio de ventana", metrics["dataset"]["evaluation_start"]),
    ], [2.0, 4.5], font_size=9.5)
    paragraph(
        doc,
        "La unidad de análisis es artista-país-fecha. Se convirtieron fecha, ranking y días en chart a tipos explícitos. La exclusión temporal evita que la falta de historia durante el primer año fuerce a clasificar a todos los artistas como nuevos. La anomalía de Venezuela ya había sido identificada en el análisis exploratorio y se mantuvo fuera para no introducir un patrón espurio.",
    )
    heading(doc, "2.2 Variables de caracterización", 2)
    table(doc, ["Variable", "Regla", "Interpretación"], [
        ("Antigüedad", ">365 días = Veterano", "Trayectoria sostenida"),
        ("Alcance", ">10 países = Global", "Difusión internacional"),
        ("Visibilidad", "rank <=50 = Top", "Posición competitiva"),
    ], [1.45, 2.0, 3.05], font_size=9.2)

    # 5 — sample
    page_break(doc)
    heading(doc, "2.3 Población, estratos y muestreo", 2)
    paragraph(
        doc,
        "Las tres variables binarias producen ocho combinaciones. La probabilidad de cada estrato se estimó empíricamente y el tamaño muestral se asignó de manera proporcional. El muestreo se realizó sin reemplazo y con semilla fija.",
    )
    rows = []
    for _, row in strata.sort_values("count_P", ascending=False).iterrows():
        rows.append((
            row["partition_id"],
            f"{int(row['count_P']):,}",
            f"{row['prop_P'] * 100:.2f}%",
            f"{int(row['count_M']):,}",
        ))
    table(doc, ["Estrato", "Población", "% P", "Muestra"], rows, [3.0, 1.2, 0.8, 1.5], font_size=8.5)
    paragraph(
        doc,
        f"La muestra final contiene {metrics['sample']['actual']:,} registros y conserva los ocho estratos. La mayor diferencia absoluta entre proporciones de población y muestra fue {metrics['sample']['max_population_sample_difference_pct']:.4f} puntos porcentuales. Esta comprobación reduce el riesgo de que el entrenamiento refleje una composición artificial.",
    )
    heading(doc, "2.4 Construcción de folds", 2)
    paragraph(
        doc,
        f"Cada estrato se ordenó de forma pseudoaleatoria y sus registros se distribuyeron cíclicamente en cinco folds. El fold menor contiene {metrics['folds']['minimum_size']:,} observaciones y el mayor {metrics['folds']['maximum_size']:,}; todos contienen los ocho estratos.",
    )

    # 6 — feature/hyperparameters
    page_break(doc)
    heading(doc, "2.5 Preparación y prevención de fuga", 2)
    paragraph(
        doc,
        "Las variables supervisadas fueron days_on_chart, max_days_on_chart, country_count_by_artist y los índices categóricos de país, antigüedad y alcance. Rank quedó excluida porque rank_tier se deriva de ella. Esta decisión es esencial: utilizar el ranking como predictor convertiría la tarea en una reconstrucción trivial de la etiqueta.",
    )
    heading(doc, "2.6 Ajuste de hiperparámetros", 2)
    paragraph(
        doc,
        "Se reservó el fold 0 para validación de cuatro configuraciones. AUC-ROC fue el criterio principal y F1 ponderado el desempate. Los pesos inversos por clase se calcularon sobre el conjunto de entrenamiento.",
    )
    tuning_rows = [
        (
            int(r.numTrees), int(r.maxDepth),
            f"{r.auc_roc:.4f}", f"{r.f1_weighted:.4f}", f"{r.accuracy:.4f}",
        )
        for r in tuning.itertuples()
    ]
    table(doc, ["Árboles", "Profundidad", "AUC-ROC", "F1", "Accuracy"], tuning_rows,
          [1.2, 1.35, 1.3, 1.3, 1.35], font_size=9.2)
    paragraph(
        doc,
        f"La configuración seleccionada fue numTrees={metrics['random_forest']['selected_numTrees']} y maxDepth={metrics['random_forest']['selected_maxDepth']}. La búsqueda es deliberadamente ligera: permite justificar complejidad sin convertir el ajuste en una exploración computacional desproporcionada.",
    )
    figure(doc, "05_importancia_variables.png", "Figura 1. Importancia promedio de variables durante la validación cruzada.", width=5.25)

    # 7 — CV results
    page_break(doc)
    heading(doc, "3. Experimentación", 1)
    heading(doc, "3.1 Validación cruzada de cinco pliegues", 2)
    paragraph(
        doc,
        "En cada iteración se entrenó con cuatro folds y se evaluó en el restante. Esta estrategia produjo cinco estimaciones independientes sobre particiones balanceadas. La variabilidad se interpreta junto con el promedio para evitar seleccionar un modelo por un resultado aislado.",
    )
    fold_rows = [
        (
            int(r.fold_id), f"{int(r.n_train):,}", f"{int(r.n_test):,}",
            f"{r.accuracy:.4f}", f"{r.f1_weighted:.4f}", f"{r.auc_roc:.4f}",
        )
        for r in folds.itertuples()
    ]
    table(doc, ["Fold", "Train", "Test", "Accuracy", "F1", "AUC-ROC"], fold_rows,
          [0.55, 1.2, 1.0, 1.2, 1.2, 1.35], font_size=8.8)
    table(doc, ["Métrica", "Promedio", "Desv. estándar"], [
        ("Accuracy", f"{metric(metrics, 'accuracy'):.4f}", f"{metric(metrics, 'accuracy', 'std'):.4f}"),
        ("F1 ponderado", f"{metric(metrics, 'f1_weighted'):.4f}", f"{metric(metrics, 'f1_weighted', 'std'):.4f}"),
        ("Precisión ponderada", f"{metric(metrics, 'precision_weighted'):.4f}", f"{metric(metrics, 'precision_weighted', 'std'):.4f}"),
        ("Recall ponderado", f"{metric(metrics, 'recall_weighted'):.4f}", f"{metric(metrics, 'recall_weighted', 'std'):.4f}"),
        ("AUC-ROC", f"{metric(metrics, 'auc_roc'):.4f}", f"{metric(metrics, 'auc_roc', 'std'):.4f}"),
    ], [2.8, 1.85, 1.85], font_size=9.5)
    baseline_rows = []
    labels = {
        "accuracy": "Accuracy",
        "f1_weighted": "F1 ponderado",
        "precision_weighted": "Precisión ponderada",
        "recall_weighted": "Recall ponderado",
        "auc_roc": "AUC-ROC",
    }
    for row in baseline.itertuples():
        model_value = metrics["random_forest"]["metrics_mean"][row.metric]
        baseline_value = metrics["random_forest"]["baseline_mean"][row.metric]
        delta_value = metrics["random_forest"]["model_minus_baseline"][row.metric]
        baseline_rows.append((
            labels[row.metric],
            f"{model_value:.4f}",
            f"{baseline_value:.4f}",
            f"{delta_value:+.4f}",
        ))
    table(doc, ["Métrica", "Random Forest", "Línea base", "Diferencia"], baseline_rows,
          [2.25, 1.45, 1.4, 1.4], font_size=8.9)
    paragraph(
        doc,
        "La línea base predice siempre Estable. Su accuracy es ligeramente mayor porque esa clase representa cerca de tres cuartas partes de la muestra, pero no identifica ningún caso Top. Random Forest sacrifica parte de esa accuracy para mejorar F1, AUC-ROC y recuperación de la clase minoritaria.",
    )

    # 8 — ROC/confusion
    heading(doc, "3.2 Capacidad de separación y errores", 2)
    paragraph(
        doc,
        f"El AUC-ROC promedio de {metric(metrics, 'auc_roc'):.4f} resume la capacidad de ordenar observaciones de las dos clases a través de distintos umbrales. Su desviación de {metric(metrics, 'auc_roc', 'std'):.4f} muestra cuánto cambia esa capacidad entre folds.",
    )
    figure(doc, "03_curvas_roc.png", "Figura 2. Curvas ROC de los cinco folds.", width=3.75)
    paragraph(
        doc,
        "La matriz de confusión del mejor fold permite localizar los errores. Debe leerse junto con F1, precisión y recall: una accuracy aceptable podría ocultar un costo desproporcionado sobre la clase minoritaria.",
    )
    figure(doc, "04_matriz_confusion.png", "Figura 3. Matriz de confusión del fold con mejor generalización.", width=3.05)
    paragraph(
        doc,
        f"El fold oficial es el {metrics['random_forest']['best_fold']}, seleccionado con una regla única: mayor AUC-ROC y F1 ponderado como desempate. Esta misma definición se utiliza en el notebook, las tablas, la figura y las conclusiones.",
        italic=True,
        color=MUTED,
    )
    class_rows = [
        (
            row["class"],
            f"{int(row['support']):,}",
            f"{row['precision']:.4f}",
            f"{row['recall']:.4f}",
            f"{row['f1']:.4f}",
        )
        for _, row in per_class.iterrows()
    ]
    table(doc, ["Clase", "Soporte", "Precisión", "Recall", "F1"], class_rows,
          [1.5, 1.25, 1.25, 1.25, 1.25], font_size=9.1)
    paragraph(
        doc,
        f"El modelo es más preciso para Estable ({metrics['random_forest']['per_class']['Estable']['precision']:.4f}) que para Top ({metrics['random_forest']['per_class']['Top']['precision']:.4f}). Aun así, recupera {metrics['random_forest']['per_class']['Top']['recall']:.2%} de los casos Top, frente a 0% en la línea base. El F1 de Top, {metrics['random_forest']['per_class']['Top']['f1']:.4f}, muestra que esta clase continúa siendo el principal espacio de mejora.",
    )

    # 9 — variability and importance
    heading(doc, "3.3 Variabilidad e interpretación", 2)
    figure(doc, "02_variabilidad_metricas.png", "Figura 4. Distribución y variabilidad de las métricas entre folds.", width=5.35)
    paragraph(
        doc,
        "La dispersión estrecha entre folds indica que el resultado no depende de una sola partición favorable. Esta estabilidad es especialmente relevante en una muestra estratificada, donde cada fold conserva perfiles raros y mayoritarios.",
    )
    imp_rows = [
        (r.feature, f"{r.mean:.4f}", f"{r.std:.4f}")
        for r in importance.itertuples()
    ]
    table(doc, ["Variable", "Importancia media", "Desv. estándar"], imp_rows,
          [3.1, 1.7, 1.7], font_size=9.2)
    paragraph(
        doc,
        "La importancia no implica causalidad. Describe la contribución relativa dentro del bosque entrenado y puede distribuirse entre variables correlacionadas. Su utilidad consiste en verificar que el modelo se apoya en señales plausibles de trayectoria, alcance y contexto.",
    )

    # 10 — kmeans
    heading(doc, "3.4 Análisis complementario con KMeans", 2)
    paragraph(
        doc,
        "El modelo no supervisado utiliza days_on_chart, max_days_on_chart, country_count_by_artist y rank estandarizados. Aquí rank puede incluirse porque no se usa una etiqueta derivada para entrenar el agrupamiento.",
    )
    k_rows = [(int(r.k), f"{r.inertia:,.0f}", f"{r.silhouette:.4f}") for r in kmeans.itertuples()]
    table(doc, ["k", "Inertia", "Silhouette"], k_rows, [1.0, 3.0, 2.5], font_size=9.5)
    figure(doc, "07_seleccion_kmeans.png", "Figura 5. Comparación de inertia y Silhouette para seleccionar k.", width=5.4)
    center_rows = []
    for r in centers.itertuples():
        center_rows.append((
            int(r.cluster), f"{r.days_on_chart:.1f}", f"{r.max_days_on_chart:.1f}",
            f"{r.country_count_by_artist:.1f}", f"{r.rank:.1f}", f"{r.pct:.1f}%",
        ))
    table(doc, ["Cluster", "Días", "Máx. días", "Países", "Rank", "%"], center_rows,
          [0.75, 1.0, 1.25, 1.1, 1.0, 1.4], font_size=8.7)
    paragraph(
        doc,
        f"Silhouette seleccionó k={metrics['kmeans']['best_k']} con valor {metrics['kmeans']['silhouette']:.4f}. Los centroides permiten describir perfiles, pero no deben confundirse con segmentos causales o identidades permanentes de los artistas.",
    )

    # 11 — discussion
    heading(doc, "4. Discusión", 1)
    paragraph(
        doc,
        "La evidencia conjunta respalda que la visibilidad competitiva está asociada con señales de permanencia, alcance geográfico y contexto. El modelo alcanza una discriminación útil sin recurrir a rank, lo que evita una conclusión circular. La consistencia entre folds fortalece la interpretación porque las métricas no dependen de una sola separación de entrenamiento y prueba.",
    )
    paragraph(
        doc,
        f"La comparación con la línea base obliga a una lectura matizada. Random Forest no mejora la accuracy mayoritaria: obtiene {metric(metrics, 'accuracy'):.4f} frente a {metrics['random_forest']['baseline_mean']['accuracy']:.4f}. Su valor está en que incrementa el F1 ponderado de {metrics['random_forest']['baseline_mean']['f1_weighted']:.4f} a {metric(metrics, 'f1_weighted'):.4f}, alcanza AUC-ROC de {metric(metrics, 'auc_roc'):.4f} y recupera casos Top que la línea base ignora por completo. El costo es una precisión menor para Top y más falsos positivos, como muestra la matriz de confusión.",
    )
    paragraph(
        doc,
        "El muestreo estratificado cumple dos funciones. Primero, reduce el costo de experimentar sobre millones de filas. Segundo, conserva perfiles minoritarios, como artistas nuevos que alcanzan el Top, que podrían perder representación en una muestra simple. La diferencia mínima entre las proporciones de población y muestra aporta evidencia cuantitativa de representatividad.",
    )
    paragraph(
        doc,
        "El análisis no supervisado muestra que la estructura observada también puede organizarse sin etiquetas. Sin embargo, la correspondencia entre clusters y estratos no es perfecta ni debería serlo: los estratos fueron definidos por reglas binarias, mientras que KMeans utiliza distancias continuas.",
    )
    heading(doc, "4.1 Limitaciones", 2)
    bullets(doc, [
        "La clasificación es explicativa y contemporánea; no mide desempeño futuro.",
        "Un artista puede aparecer en más de un fold, por lo que no se prueba generalización a artistas completamente nuevos.",
        "El índice categórico de país no expresa cercanía cultural o regional.",
        "El muestreo disminuye costo, pero deja fuera parte de la población.",
        "La importancia de variables del bosque no equivale a efecto causal.",
    ])

    # 12 — conclusions and links
    heading(doc, "5. Conclusiones y trabajo futuro", 1)
    paragraph(
        doc,
        "El proyecto integró de forma reproducible las etapas solicitadas: lectura del dataset, auditoría, caracterización, construcción de muestra, preparación, ajuste de hiperparámetros, entrenamiento, validación, visualización y discusión. La muestra preservó ocho estratos y los cinco folds mantuvieron tamaños prácticamente idénticos.",
    )
    paragraph(
        doc,
        f"Random Forest alcanzó AUC-ROC promedio de {metric(metrics, 'auc_roc'):.4f} y F1 ponderado de {metric(metrics, 'f1_weighted'):.4f}. No superó la accuracy de la línea base mayoritaria, pero elevó el F1 ponderado y recuperó {metrics['random_forest']['per_class']['Top']['recall']:.2%} de la clase Top, que la línea base nunca detecta. La baja variabilidad permite concluir que ese intercambio es estable dentro del diseño experimental. KMeans añadió una lectura de perfiles y alcanzó Silhouette de {metrics['kmeans']['silhouette']:.4f} con k={metrics['kmeans']['best_k']}.",
    )
    heading(doc, "5.1 Trabajo futuro", 2)
    bullets(doc, [
        "Usar división temporal para evaluar comportamiento sobre fechas posteriores.",
        "Asignar artistas completos a un solo fold para medir generalización a entidades no vistas.",
        "Comparar Random Forest con Gradient-Boosted Trees y regresión logística.",
        "Crear variables históricas calculadas únicamente con información anterior a cada observación.",
        "Ejecutar el flujo sobre infraestructura distribuida y la población completa.",
    ])
    heading(doc, "6. Enlaces de entrega", 1)
    paragraph(doc, f"Notebook Jupyter: {NOTEBOOK_LINK}", bold=True, color=DARK_BLUE, align=WD_ALIGN_PARAGRAPH.LEFT)
    paragraph(doc, f"Video explicativo: {VIDEO_LINK}", bold=True, color=DARK_BLUE, align=WD_ALIGN_PARAGRAPH.LEFT)
    paragraph(
        doc,
        "Estos marcadores deben reemplazarse por enlaces públicos o compartidos antes de cargar el documento final en la plataforma.",
        italic=True, color=MUTED, align=WD_ALIGN_PARAGRAPH.LEFT,
    )

    # 13 — references
    heading(doc, "Referencias", 1)
    references = [
        "Ahmed, S. K. (2024). How to choose a sampling technique and determine sample size for research: A simplified guide for researchers. Oral Oncology Reports, 12, 100662. https://doi.org/10.1016/j.oor.2024.100662",
        "Apache Software Foundation. (2025). MLlib: Main guide. https://spark.apache.org/docs/latest/ml-guide.html",
        "Breiman, L. (2001). Random forests. Machine Learning, 45, 5-32. https://doi.org/10.1023/A:1010933404324",
        "International Federation of the Phonographic Industry. (2025). Global music report 2025: State of the industry. https://www.ifpi.org/",
        "Kim, J. K., & Wang, Z. (2019). Sampling techniques for big data analysis. International Statistical Review, 87(S1), S177-S191. https://doi.org/10.1111/insr.12290",
        "Lakens, D. (2022). Sample size justification. Collabra: Psychology, 8(1). https://doi.org/10.1525/collabra.33267",
        "Lopez Gil, G. (s. f.). Spotify Charts Daily Updated [Conjunto de datos]. Kaggle. https://www.kaggle.com/datasets/gonzalopezgil/spotify-charts-daily-updated",
        "Rousseeuw, P. J. (1987). Silhouettes: A graphical aid to the interpretation and validation of cluster analysis. Journal of Computational and Applied Mathematics, 20, 53-65. https://doi.org/10.1016/0377-0427(87)90125-7",
        "Spotify. (2025, 12 de marzo). How the music industry's cultural and financial impact define its success in 2025. https://newsroom.spotify.com/",
    ]
    for ref in references:
        p = paragraph(doc, ref, size=10.5, align=WD_ALIGN_PARAGRAPH.LEFT, after=7, line=1.2)
        p.paragraph_format.left_indent = Inches(0.35)
        p.paragraph_format.first_line_indent = Inches(-0.35)

    doc.save(OUTPUT)
    enforce_arial_12(OUTPUT)
    print(f"Creado: {OUTPUT}")


if __name__ == "__main__":
    build()
